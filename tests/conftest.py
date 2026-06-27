from pathlib import Path

import fitz
import pytest
from docx import Document
from PIL import Image


@pytest.fixture
def docx_template(tmp_path: Path) -> Path:
    path = tmp_path / "assignment.docx"
    document = Document()
    document.sections[0].header.paragraphs[0].text = "程序设计课程作业"
    document.sections[0].footer.paragraphs[0].text = "请保留本页格式"
    document.add_heading("第1题", level=1)
    document.add_paragraph("编写一个输出问候语的程序。")
    document.add_paragraph("代码：")
    document.add_paragraph("")
    document.add_paragraph("运行结果：")
    document.add_paragraph("")
    document.add_heading("2、第二题", level=1)
    document.add_paragraph("程序：")
    document.add_paragraph("在此填写代码")
    document.add_paragraph("截图：")
    document.add_paragraph("")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "教师评分栏"
    document.save(path)
    return path


@pytest.fixture
def unsafe_docx_template(tmp_path: Path) -> Path:
    path = tmp_path / "unsafe.docx"
    document = Document()
    document.add_heading("第1题", level=1)
    document.add_paragraph("代码：")
    document.add_paragraph("教师示例不能删除")
    document.add_paragraph("运行结果：")
    document.add_paragraph("")
    document.save(path)
    return path


@pytest.fixture
def answer_dirs(tmp_path: Path) -> tuple[Path, Path]:
    answers = tmp_path / "answers"
    screenshots = tmp_path / "screenshots"
    answers.mkdir()
    screenshots.mkdir()
    (answers / "q1.py").write_text('print("hello")\n', encoding="utf-8")
    Image.new("RGB", (640, 360), "white").save(screenshots / "q1-1.png")
    return answers, screenshots


@pytest.fixture
def pdf_template(tmp_path: Path) -> Path:
    path = tmp_path / "assignment.pdf"
    document = fitz.open()
    page = document.new_page(width=595, height=842)
    page.insert_font(fontname="china-s")
    page.insert_text((72, 72), "第1题", fontname="china-s", fontsize=14)
    page.insert_text((72, 110), "代码：", fontname="china-s", fontsize=12)
    page.insert_text((72, 360), "运行结果：", fontname="china-s", fontsize=12)
    document.save(path)
    document.close()
    return path

